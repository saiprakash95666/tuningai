"""
Resume Service - Parse and generate resumes
Handles DOCX and PDF formats
"""

from docx import Document
from PyPDF2 import PdfReader
from typing import Optional, BinaryIO
import io
import re


class ResumeService:
    """Handle resume parsing and generation"""
    
    async def parse_resume(self, file: BinaryIO, filename: str) -> str:
        """
        Parse resume from uploaded file
        
        Args:
            file: File object
            filename: Original filename
            
        Returns:
            Extracted text from resume
        """
        
        # Read file content
        content = await file.read()
        
        # Parse based on file type
        if filename.lower().endswith('.docx'):
            return self._parse_docx(content)
        elif filename.lower().endswith('.pdf'):
            return self._parse_pdf(content)
        else:
            raise ValueError(f"Unsupported file format: {filename}. Use DOCX or PDF.")
    
    def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX file"""
        try:
            doc = Document(io.BytesIO(content))
            
            # Extract all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables (many resumes use tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            resume_text = "\n".join(text_parts)
            
            if not resume_text.strip():
                raise ValueError("DOCX file appears to be empty")
            
            return resume_text
            
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def _parse_pdf(self, content: bytes) -> str:
        """Parse PDF file"""
        try:
            pdf_reader = PdfReader(io.BytesIO(content))
            
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            resume_text = "\n".join(text_parts)
            
            if not resume_text.strip():
                raise ValueError("PDF file appears to be empty or unreadable")
            
            return resume_text
            
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def extract_sections(self, resume_text: str) -> dict:
        """
        Extract common resume sections
        Used for targeted improvements
        """
        sections = {
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "projects": "",
            "other": ""
        }
        
        # Simple section detection (can be improved)
        current_section = "other"
        lines = resume_text.split("\n")
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Detect section headers
            if any(keyword in line_lower for keyword in ["summary", "objective", "profile"]):
                current_section = "summary"
            elif any(keyword in line_lower for keyword in ["experience", "employment", "work history"]):
                current_section = "experience"
            elif "education" in line_lower:
                current_section = "education"
            elif "skill" in line_lower:
                current_section = "skills"
            elif "project" in line_lower:
                current_section = "projects"
            
            # Add line to current section
            if line.strip():
                sections[current_section] += line + "\n"
        
        return sections
    
    def apply_improvements(
        self,
        resume_text: str,
        improvements: list
    ) -> str:
        """
        Apply suggested improvements to resume text
        Simple text replacement for now
        """
        improved_text = resume_text
        
        for improvement in improvements:
            current = improvement.get('current', '')
            new = improvement.get('improved', '')
            
            if current and current != 'MISSING':
                # Replace current text with improved version
                improved_text = improved_text.replace(current, new)
            else:
                # Add new section at the end
                section = improvement.get('section', 'Additional')
                improved_text += f"\n\n{section}:\n{new}"
        
        return improved_text
    
    def generate_docx(
        self,
        resume_text: str,
        output_path: str = "resume_optimized.docx"
    ) -> str:
        """
        Generate DOCX file from text
        Returns path to generated file
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading('Resume', 0)
        title.alignment = 1  # Center alignment
        
        # Add content
        for line in resume_text.split('\n'):
            if line.strip():
                # Detect if line is a header (all caps or starts with #)
                if line.isupper() or line.startswith('#'):
                    doc.add_heading(line.strip('#').strip(), level=1)
                else:
                    doc.add_paragraph(line)
        
        # Save
        import tempfile
        import os
        
        # Save to temp directory
        temp_dir = tempfile.gettempdir()
        full_path = os.path.join(temp_dir, output_path)
        doc.save(full_path)
        
        return full_path


# Singleton
_resume_service: Optional[ResumeService] = None


def get_resume_service() -> ResumeService:
    """Get resume service instance"""
    global _resume_service
    if _resume_service is None:
        _resume_service = ResumeService()
    return _resume_service